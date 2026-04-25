import json
import os
import re
import requests
import tkinter as tk
from pathlib import Path
from tkinter import filedialog, messagebox
from tkinter.scrolledtext import ScrolledText

from docx import Document
from openpyxl import Workbook, load_workbook

OLLAMA_API = "http://127.0.0.1:11434/v1/completions"
MODEL_NAME = "deepseek"
BACKEND_OPTIONS = ["Ollama", "OpenClaw"]


def parse_ollama_response(payload: dict) -> str:
    if "choices" in payload and payload["choices"]:
        message = payload["choices"][0].get("message") or payload["choices"][0]
        if isinstance(message, dict):
            return message.get("content", "")
        return str(message)

    if "result" in payload:
        return str(payload["result"])

    if "output" in payload:
        return str(payload["output"])

    return json.dumps(payload, indent=2, ensure_ascii=False)


def request_ollama(url: str, json_data: dict) -> str:
    response = requests.post(url, json=json_data, timeout=30)
    response.raise_for_status()
    return parse_ollama_response(response.json())


def ollama_query(prompt: str) -> str:
    api_url = api_url_var.get().strip() or OLLAMA_API
    base_url = api_url.rstrip("/")

    chat_url = base_url + "/v1/chat/completions"
    completion_url = base_url + "/v1/completions"

    payload_chat = {
        "model": MODEL_NAME,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.2,
    }
    payload_completions = {
        "model": MODEL_NAME,
        "prompt": prompt,
        "temperature": 0.2,
    }

    if api_url.endswith("/v1/chat/completions"):
        return request_ollama(api_url, payload_chat)
    if api_url.endswith("/v1/completions"):
        return request_ollama(api_url, payload_completions)

    try:
        return request_ollama(completion_url, payload_completions)
    except requests.HTTPError as exc:
        if exc.response is not None and exc.response.status_code == 404:
            return request_ollama(chat_url, payload_chat)
        raise


def openclaw_query(prompt: str) -> str:
    try:
        from openclaw import OpenClaw
    except ImportError as exc:
        raise RuntimeError(
            "Không thể import OpenClaw. Hãy cài openclaw trong virtualenv bằng `pip install openclaw` và kiểm tra dependencies."
        ) from exc

    try:
        client = OpenClaw.local()
    except Exception as exc:
        raise RuntimeError(
            "Không kết nối được OpenClaw local. Hãy kiểm tra CMDOP agent đang chạy hoặc cài đặt OpenClaw đúng cách."
        ) from exc

    result = client.agent.run(prompt)
    if hasattr(result, "text"):
        return result.text
    return str(result)


def save_word(file_path: Path, text: str) -> None:
    document = Document(file_path) if file_path.exists() else Document()
    document.add_paragraph(text)
    document.save(file_path)


def save_excel(file_path: Path, text: str) -> None:
    workbook = load_workbook(file_path) if file_path.exists() else Workbook()
    sheet = workbook.active

    # Xóa dữ liệu cũ nếu file đã tồn tại
    if file_path.exists():
        sheet.delete_rows(1, sheet.max_row)

    # Thêm headers
    headers = ["STT", "Nội dung", "Mô tả", "Ghi chú"]
    for col_num, header in enumerate(headers, 1):
        sheet.cell(row=1, column=col_num, value=header)

    # Parse nội dung thành rows
    rows = []
    lines = text.strip().splitlines()
    for idx, line in enumerate(lines, start=1):
        # Thử parse theo tab hoặc pipe
        if "\t" in line:
            cells = [cell.strip() for cell in line.split("\t")]
        elif "|" in line:
            cells = [cell.strip() for cell in line.split("|") if cell.strip()]
        else:
            # Nếu không có dấu phân cách, chia thành các phần nhỏ hơn
            cells = [line.strip()]

        # Đảm bảo có đủ cột
        while len(cells) < len(headers) - 1:  # trừ STT
            cells.append("")

        # Thêm STT và cells
        row = [idx] + cells[:len(headers)-1]
        rows.append(row)

    # Ghi rows vào sheet
    for r, row in enumerate(rows, start=2):  # bắt đầu từ row 2
        for c, value in enumerate(row, start=1):
            sheet.cell(row=r, column=c, value=value)

    workbook.save(file_path)


def create_excel_file() -> None:
    if file_type_var.get() != "Excel":
        messagebox.showwarning("Chọn Excel", "Vui lòng chọn loại file Excel để tạo file mới.")
        return

    output_path = file_path_var.get().strip()
    if not output_path:
        messagebox.showwarning("Thiếu đường dẫn", "Vui lòng chọn vị trí và tên file Excel để tạo.")
        return

    save_path = Path(output_path)
    if save_path.exists():
        overwrite = messagebox.askyesno(
            "File đã tồn tại",
            f"File {save_path.name} đã tồn tại. Bạn có muốn tạo lại file mới và ghi đè?",
        )
        if not overwrite:
            status_var.set("Đã hủy tạo file Excel.")
            return

    workbook = Workbook()
    workbook.save(save_path)
    status_var.set(f"Đã tạo file Excel mới: {save_path}")
    messagebox.showinfo("Hoàn thành", f"Đã tạo file Excel: {save_path}")


def choose_file() -> None:
    file_type = file_type_var.get()
    if file_type == "Word":
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            title="Chọn file Word",
        )
    else:
        path = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel Workbook", "*.xlsx")],
            title="Chọn file Excel",
        )
    if path:
        file_path_var.set(path)


def run_action() -> None:
    prompt = prompt_text.get("1.0", tk.END).strip()
    output_path = file_path_var.get().strip()
    backend = backend_var.get()
    if not prompt:
        messagebox.showwarning("Thiếu thông tin", "Vui lòng nhập yêu cầu.")
        return
    if not output_path:
        messagebox.showwarning("Thiếu thông tin", "Vui lòng chọn file Word hoặc Excel.")
        return

    try:
        status_var.set(f"Đang gửi yêu cầu tới {backend}...")
        root.update_idletasks()
        if backend == "OpenClaw":
            result = openclaw_query(prompt)
        else:
            result = ollama_query(prompt)

        save_path = Path(output_path)
        if file_type_var.get() == "Word":
            save_word(save_path, result)
        else:
            save_excel(save_path, result)

        output_text.delete("1.0", tk.END)
        output_text.insert(tk.END, result)
        status_var.set(f"Hoàn thành: đã lưu vào {save_path}")
    except requests.RequestException as exc:
        messagebox.showerror("Lỗi kết nối Ollama", f"Không thể liên hệ tới Ollama:\n{exc}")
        status_var.set("Lỗi: kiểm tra Ollama đang chạy và API endpoint.")
    except Exception as exc:
        messagebox.showerror("Lỗi", str(exc))
        status_var.set(str(exc))


root = tk.Tk()
root.title("Ollama Deepseek - Word/Excel Local UI")
root.geometry("760x640")
root.resizable(False, False)

file_type_var = tk.StringVar(value="Word")
file_path_var = tk.StringVar(value="")
status_var = tk.StringVar(value="Chưa kết nối.")

frame_top = tk.Frame(root, padx=12, pady=12)
frame_top.pack(fill=tk.X)

api_url_var = tk.StringVar(value=OLLAMA_API)
backend_var = tk.StringVar(value=BACKEND_OPTIONS[0])
label_backend = tk.Label(frame_top, text="Backend:", anchor="w")
label_backend.grid(row=0, column=0, sticky="w")
option_backend = tk.OptionMenu(frame_top, backend_var, *BACKEND_OPTIONS)
option_backend.grid(row=0, column=1, sticky="w")

label_api = tk.Label(frame_top, text="API Ollama (Ollama only):", anchor="w")
label_api.grid(row=1, column=0, sticky="w")
entry_api = tk.Entry(frame_top, width=60, textvariable=api_url_var)
entry_api.grid(row=1, column=1, columnspan=3, sticky="w", padx=(8, 0))

label_model = tk.Label(frame_top, text="Model:", anchor="w")
label_model.grid(row=2, column=0, sticky="w", pady=(8, 0))
entry_model = tk.Entry(frame_top, width=20)
entry_model.insert(0, MODEL_NAME)
entry_model.config(state="readonly")
entry_model.grid(row=2, column=1, sticky="w", pady=(8, 0), padx=(8, 0))

label_type = tk.Label(frame_top, text="Loại file:", anchor="w")
label_type.grid(row=3, column=0, sticky="w", pady=(8, 0))
option_menu = tk.OptionMenu(frame_top, file_type_var, "Word", "Excel")
option_menu.grid(row=3, column=1, sticky="w", pady=(8, 0), padx=(8, 0))

label_path = tk.Label(frame_top, text="File output:", anchor="w")
label_path.grid(row=4, column=0, sticky="w", pady=(8, 0))
entry_path = tk.Entry(frame_top, textvariable=file_path_var, width=50)
entry_path.grid(row=4, column=1, sticky="w", pady=(8, 0), padx=(8, 0))
button_path = tk.Button(frame_top, text="Chọn file…", command=choose_file)
button_path.grid(row=4, column=2, sticky="w", pady=(8, 0), padx=(8, 0))

label_prompt = tk.Label(root, text="Nhập yêu cầu:")
label_prompt.pack(anchor="w", padx=12)
prompt_text = ScrolledText(root, width=90, height=10)
prompt_text.pack(padx=12, pady=(0, 12))

button_frame = tk.Frame(root, padx=12)
button_frame.pack(fill=tk.X)
button_run = tk.Button(button_frame, text="Chạy và lưu kết quả", command=run_action, width=24)
button_run.pack(side=tk.LEFT)
button_create_excel = tk.Button(button_frame, text="Tạo file Excel", command=create_excel_file, width=18)
button_create_excel.pack(side=tk.LEFT, padx=(8, 0))

status_label = tk.Label(root, textvariable=status_var, anchor="w", fg="#084B8A")
status_label.pack(fill=tk.X, padx=12, pady=(8, 0))

label_output = tk.Label(root, text="Kết quả trả về:")
label_output.pack(anchor="w", padx=12, pady=(12, 0))
output_text = ScrolledText(root, width=90, height=16)
output_text.pack(padx=12, pady=(0, 12))

root.mainloop()
